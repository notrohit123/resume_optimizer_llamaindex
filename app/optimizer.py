import os
import streamlit as st
import tempfile
import base64
import unicodedata

from llama_index.llms.openai import OpenAI
from llama_index.llms.anthropic import Anthropic
from llama_index.embeddings.openai import OpenAIEmbedding
from llama_index.core import Settings, VectorStoreIndex, Document

from docx import Document as DocxDocument
from fpdf import FPDF


def sanitize_text(text):
    return unicodedata.normalize("NFKD", text).encode("latin-1", "replace").decode("latin-1")

def get_llm(gen_model, api_key=None):
    if api_key is None:
        api_key = os.getenv("OPENAI_API_KEY")

    if gen_model.startswith("gpt-"):
        return OpenAI(model=gen_model, api_key=api_key)
    elif gen_model.startswith("claude-"):
        return Anthropic(model=gen_model, api_key=os.getenv("ANTHROPIC_API_KEY"))
    else:
        raise ValueError(f"Unsupported model: {gen_model}")


def optimize_resume(text, prompt, job_title, job_description, gen_model, api_key):
    if not api_key:
        raise ValueError("OpenAI API key is missing!")

    embed_model = OpenAIEmbedding(model="text-embedding-3-small", api_key=api_key)
    llm = get_llm(gen_model, api_key=api_key)

    Settings.llm = llm
    Settings.embed_model = embed_model

    doc = Document(text=text)
    index = VectorStoreIndex.from_documents([doc])

    resume_analysis = index.as_query_engine(similarity_top_k=1).query(
        "Analyze this resume and highlight key strengths and gaps."
    )

    full_prompt = f"""
    Resume Analysis:
    {resume_analysis}

    Job Title: {job_title}
    Job Description: {job_description}

    Optimization Request: {prompt}

    Provide:
    ## Key Findings
    • Summary of alignment and gaps

    ## Specific Improvements
    • Bullet points with actionable changes
    • Use strong verbs and relevant examples

    ## Actionable Suggestions
    • Immediate next steps
    """

    result = index.as_query_engine(similarity_top_k=1).query(full_prompt)
    return result.response


def get_ats_score(resume_text, job_description):
    job_keywords = set(job_description.lower().split())
    resume_words = set(resume_text.lower().split())
    match = job_keywords & resume_words
    score = (len(match) / len(job_keywords)) * 100
    return round(score, 2), list(match)


def generate_tailored_resume(text, job_title, job_description, model):
    llm = get_llm(model)
    prompt = f"""
    You are an expert resume writer. Rewrite this resume so that it is perfectly tailored to the following job title and description:

    Job Title: {job_title}
    Job Description: {job_description}

    Resume to Rewrite:
    {text}

    Output only the new tailored resume.
    """

    response = llm.complete(prompt).text

   
    st.download_button("⬇️ Download as TXT", response, file_name="tailored_resume.txt")

   
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    for line in sanitize_text(response).split("\n"):
        pdf.multi_cell(0, 10, line)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
        pdf.output(tmp_pdf.name)
        with open(tmp_pdf.name, "rb") as f:
            b64_pdf = base64.b64encode(f.read()).decode()
            href = f'<a href="data:application/pdf;base64,{b64_pdf}" download="tailored_resume.pdf">⬇️ Download as PDF</a>'
            st.markdown(href, unsafe_allow_html=True)

  
    doc = DocxDocument()
    for line in response.split("\n"):
        doc.add_paragraph(line)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
        doc.save(tmp_docx.name)
        with open(tmp_docx.name, "rb") as f:
            b64_docx = base64.b64encode(f.read()).decode()
            href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64_docx}" download="tailored_resume.docx">⬇️ Download as Word</a>'
            st.markdown(href, unsafe_allow_html=True)

    return response
